"""
Word文書処理モジュール
単一責任: Word文書の編集処理（先頭行削除）
"""
import os
import shutil
import zipfile
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Optional, List
from docx import Document
import logging

from ..core.models import ProcessingResult, LogLevel
from ..core.exceptions import WordProcessingError
from ..utils.validators import FileValidator


logger = logging.getLogger(__name__)


class WordProcessor:
    """Word文書処理クラス"""
    
    def __init__(self):
        """初期化"""
        self.validator = FileValidator()
        self._stop_requested = False
        
    def stop(self) -> None:
        """処理停止要求"""
        self._stop_requested = True
        
    def _check_stop(self) -> None:
        """停止要求をチェック"""
        if self._stop_requested:
            raise WordProcessingError("処理が中断されました")
            
    def process_docx_file(self, docx_path: str) -> bool:
        """
        Word文書の先頭行を削除
        
        Args:
            docx_path: Word文書のパス
            
        Returns:
            bool: 処理成功フラグ
        """
        try:
            doc = Document(docx_path)
            
            if not doc.paragraphs:
                logger.warning(f"空の文書: {docx_path}")
                return False
                
            # 先頭段落を削除
            first_paragraph = doc.paragraphs[0]
            p_element = first_paragraph._element
            p_element.getparent().remove(p_element)
            
            # 保存
            doc.save(docx_path)
            logger.info(f"先頭行削除完了: {docx_path}")
            return True
            
        except Exception as e:
            logger.error(f"Word文書処理エラー: {docx_path} - {str(e)}")
            return False
            
    def process_zip_file(self, zip_path: str) -> ProcessingResult:
        """
        ZIPファイル内の全Word文書を処理
        
        Args:
            zip_path: ZIPファイルのパス
            
        Returns:
            ProcessingResult: 処理結果
        """
        result = ProcessingResult(success=True)
        
        # 入力検証
        if not self.validator.validate_zip_file(zip_path):
            result.success = False
            result.add_message("無効なZIPファイル", LogLevel.ERROR)
            return result
            
        try:
            # 新しいZIPファイル名を生成
            new_zip_path = self._generate_output_filename(zip_path)
            
            with TemporaryDirectory() as temp_dir:
                # ZIP展開
                self._extract_zip(zip_path, temp_dir, result)
                
                # Word文書処理
                self._process_docx_files(temp_dir, result)
                
                # 新しいZIPに圧縮
                self._create_zip(temp_dir, new_zip_path, result)
                
                result.details['output_path'] = new_zip_path
                result.add_message(f"処理完了: {os.path.basename(new_zip_path)}", LogLevel.INFO)
                
        except Exception as e:
            logger.exception("ZIP処理エラー")
            result.success = False
            result.add_message(str(e), LogLevel.ERROR)
            
        return result
        
    def _generate_output_filename(self, original_path: str) -> str:
        """出力ファイル名を生成"""
        base_path = Path(original_path)
        base_name = base_path.stem
        directory = base_path.parent
        
        # 基本の出力名
        new_name = f"{base_name}-1linedel.zip"
        new_path = directory / new_name
        
        # 重複チェック
        counter = 1
        while new_path.exists():
            new_name = f"{base_name}-1linedel({counter}).zip"
            new_path = directory / new_name
            counter += 1
            
        return str(new_path)
        
    def _extract_zip(self, zip_path: str, extract_to: str, result: ProcessingResult) -> None:
        """ZIPファイルを展開"""
        try:
            with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                zip_ref.extractall(extract_to)
                result.add_message("ZIP展開完了", LogLevel.INFO)
        except Exception as e:
            raise WordProcessingError(f"ZIP展開エラー: {str(e)}")
            
    def _process_docx_files(self, directory: str, result: ProcessingResult) -> None:
        """ディレクトリ内のWord文書を処理"""
        docx_files = Path(directory).rglob("*.docx")
        
        for docx_file in docx_files:
            self._check_stop()
            
            if self.process_docx_file(str(docx_file)):
                result.processed_count += 1
                result.add_message(f"処理完了: {docx_file.name}", LogLevel.INFO)
            else:
                result.warning_count += 1
                result.add_message(f"処理スキップ: {docx_file.name}", LogLevel.WARNING)
                
    def _create_zip(self, source_dir: str, output_path: str, result: ProcessingResult) -> None:
        """新しいZIPファイルを作成"""
        try:
            with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
                source_path = Path(source_dir)
                
                for file_path in source_path.rglob("*"):
                    if file_path.is_file():
                        arcname = file_path.relative_to(source_path)
                        zip_ref.write(file_path, arcname)
                        
            result.add_message(f"ZIP作成完了: {os.path.basename(output_path)}", LogLevel.INFO)
            
        except Exception as e:
            raise WordProcessingError(f"ZIP作成エラー: {str(e)}")