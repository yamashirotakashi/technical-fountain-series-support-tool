from __future__ import annotations
"""Word文書処理モジュール"""
import zipfile
import tempfile
from pathlib import Path
from typing import List, Optional
import docx

from utils.logger import get_logger
from core.di_container import inject
from core.configuration_provider import ConfigurationProvider


class WordProcessor:
    """Word文書を処理するクラス"""
    
    @inject
    def __init__(self, config_provider: ConfigurationProvider):
        """
        WordProcessorを初期化
        
        Args:
            config_provider: 統一設定プロバイダー（DI注入）
        """
        self.logger = get_logger(__name__)
        self.config_provider = config_provider
    
    def process_word_files(self, folder_path: Path) -> int:
        """
        フォルダ内のすべてのWordファイルを処理
        
        Args:
            folder_path: 処理するフォルダのパス
        
        Returns:
            処理したファイル数
        """
        self.logger.info(f"Word処理開始: {folder_path}")
        
        if not folder_path.exists():
            self.logger.error(f"フォルダが存在しません: {folder_path}")
            return 0
        
        # .docxファイルを検索
        word_files = list(folder_path.glob("**/*.docx"))
        
        if not word_files:
            self.logger.warning("Wordファイルが見つかりませんでした")
            return 0
        
        processed_count = 0
        for word_file in word_files:
            try:
                if self.remove_first_line(word_file):
                    processed_count += 1
            except Exception as e:
                self.logger.error(f"ファイル処理エラー {word_file}: {e}")
        
        self.logger.info(f"Word処理完了: {processed_count}/{len(word_files)} ファイル")
        return processed_count
    
    def remove_first_line(self, doc_path: Path) -> bool:
        """
        Word文書の1行目を削除
        
        Args:
            doc_path: 処理するWord文書のパス
        
        Returns:
            処理が成功した場合True
        """
        try:
            self.logger.info(f"1行目削除処理: {doc_path}")
            
            # ドキュメントを開く
            doc = docx.Document(doc_path)
            
            # パラグラフが存在するか確認
            if not doc.paragraphs:
                self.logger.warning(f"空のドキュメント: {doc_path}")
                return True
            
            # 最初のパラグラフを削除
            first_para = doc.paragraphs[0]
            p = first_para._element
            p.getparent().remove(p)
            
            # 保存
            doc.save(doc_path)
            
            self.logger.info(f"1行目を削除しました: {doc_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"1行目削除エラー {doc_path}: {e}")
            return False
    
    def get_word_files(self, folder_path: Path) -> List[Path]:
        """
        フォルダ内のWordファイルのリストを取得
        
        Args:
            folder_path: 検索するフォルダのパス
        
        Returns:
            Wordファイルのパスリスト
        """
        if not folder_path.exists():
            return []
        
        # すべてのファイルを列挙（デバッグ用）
        all_files = list(folder_path.rglob("*"))
        self.logger.info(f"フォルダ内の全ファイル数: {len(all_files)}")
        for file_path in all_files:
            if file_path.is_file():
                self.logger.debug(f"  - {file_path.relative_to(folder_path)}")
        
        # .docxファイルを検索
        word_files = list(folder_path.glob("**/*.docx"))
        self.logger.info(f"見つかったWordファイル数: {len(word_files)}")
        
        return word_files
    
    def process_zip_file(self, zip_path: Path, temp_dir: Optional[Path] = None) -> List[Path]:
        """
        ZIPファイルからWordファイルを抽出し、1行目を削除
        
        Args:
            zip_path: 処理するZIPファイルのパス
            temp_dir: 一時ディレクトリ（指定しない場合は自動作成）
        
        Returns:
            処理済みWordファイルのパスリスト
        """
        self.logger.info(f"ZIP処理開始: {zip_path}")
        
        if not zip_path.exists():
            self.logger.error(f"ZIPファイルが存在しません: {zip_path}")
            return []
        
        # 一時ディレクトリの準備
        if temp_dir is None:
            temp_dir = Path(tempfile.mkdtemp())
        else:
            temp_dir.mkdir(parents=True, exist_ok=True)
        
        processed_files = []
        
        try:
            # ZIPファイルを展開
            with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                # ZIP内のファイル一覧をログ出力
                file_list = zip_ref.namelist()
                self.logger.info(f"ZIP内のファイル数: {len(file_list)}")
                for file_name in file_list:
                    self.logger.info(f"  - {file_name}")
                
                zip_ref.extractall(temp_dir)
            
            self.logger.info(f"ZIP展開完了: {temp_dir}")
            
            # 展開されたWordファイルを検索
            word_files = self.get_word_files(temp_dir)
            
            if not word_files:
                self.logger.warning("Wordファイルが見つかりませんでした")
                return []
            
            self.logger.info(f"Wordファイル {len(word_files)}個を発見")
            
            # 各Wordファイルの1行目を削除
            for word_file in word_files:
                try:
                    if self.remove_first_line(word_file):
                        processed_files.append(word_file)
                        self.logger.info(f"処理完了: {word_file.name}")
                    else:
                        self.logger.error(f"処理失敗: {word_file.name}")
                except Exception as e:
                    self.logger.error(f"ファイル処理エラー {word_file}: {e}")
            
            self.logger.info(f"ZIP処理完了: {len(processed_files)}/{len(word_files)} ファイル")
            return processed_files
            
        except zipfile.BadZipFile:
            self.logger.error(f"無効なZIPファイル: {zip_path}")
            return []
        except Exception as e:
            self.logger.error(f"ZIP処理エラー: {e}")
            return []
    
    def find_ncode_folder(self, ncode: str) -> Optional[Path]:
        """
        指定されたNコードのフォルダを検索
        
        Args:
            ncode: 検索するNコード
        
        Returns:
            見つかったNフォルダのパス、見つからない場合はNone
        """
        # ConfigManagerから設定を取得
        if self.config_manager:
            try:
                base_path_str = self.config_manager.get('paths.ncode_base_path')
                if base_path_str:
                    base_path = Path(base_path_str)
                    self.logger.info(f"ConfigManagerからNコードベースパスを取得: {base_path}")
                else:
                    # 環境変数からのフォールバック
                    env_path = os.environ.get('TECHZIP_NCODE_BASE_PATH')
                    if env_path:
                        base_path = Path(env_path)
                        self.logger.info(f"環境変数からNコードベースパスを取得: {base_path}")
                    else:
                        raise ValueError("Nコードベースパスが設定されていません")
            except Exception as e:
                self.logger.error(f"ConfigManagerからの設定取得に失敗: {e}")
                # 環境変数からの最終フォールバック
                env_path = os.environ.get('TECHZIP_NCODE_BASE_PATH')
                if env_path:
                    base_path = Path(env_path)
                    self.logger.warning(f"環境変数からフォールバック: {base_path}")
                else:
                    raise ValueError(
                        "Nコードベースパスが設定されていません。"
                        "ConfigManagerまたは環境変数TECHZIP_NCODE_BASE_PATHを設定してください。"
                    )
        else:
            # ConfigManagerが利用できない場合は環境変数を確認
            env_path = os.environ.get('TECHZIP_NCODE_BASE_PATH')
            if env_path:
                base_path = Path(env_path)
                self.logger.info(f"ConfigManager未利用、環境変数から取得: {base_path}")
            else:
                raise ValueError(
                    "ConfigManagerが未初期化で、環境変数TECHZIP_NCODE_BASE_PATHも設定されていません。"
                    "設定を確認してください。"
                )
        
        ncode_folder = base_path / ncode
        
        if ncode_folder.exists() and ncode_folder.is_dir():
            self.logger.info(f"Nフォルダ発見: {ncode_folder}")
            return ncode_folder
        else:
            self.logger.warning(f"Nフォルダが見つかりません: {ncode_folder}")
            return None
    
    def find_honbun_folder(self, ncode_folder: Path) -> Optional[Path]:
        """
        Nフォルダ配下の「本文」フォルダを検索
        
        Args:
            ncode_folder: Nフォルダのパス
        
        Returns:
            見つかった本文フォルダのパス、見つからない場合はNone
        """
        # ConfigManagerから設定を取得、失敗時はデフォルト値を使用
        if self.config_manager:
            try:
                honbun_folder_name = self.config_manager.get('folders.honbun_folder_name')
                if honbun_folder_name:
                    self.logger.info(f"ConfigManagerから本文フォルダ名を取得: {honbun_folder_name}")
                else:
                    # 設定が空の場合はデフォルト値
                    honbun_folder_name = "本文"
                    self.logger.info("ConfigManagerの設定が空のため、デフォルト本文フォルダ名を使用")
            except Exception as e:
                self.logger.warning(f"ConfigManagerからの設定取得に失敗: {e}")
                honbun_folder_name = "本文"
        else:
            # ConfigManagerが利用できない場合はデフォルト値
            honbun_folder_name = "本文"
            self.logger.info("ConfigManager未利用のため、デフォルト本文フォルダ名を使用")
        
        honbun_folder = ncode_folder / honbun_folder_name
        
        if honbun_folder.exists() and honbun_folder.is_dir():
            self.logger.info(f"本文フォルダ発見: {honbun_folder}")
            return honbun_folder
        else:
            # 本文フォルダが存在しない場合は作成可能かチェック
            if ncode_folder.exists():
                self.logger.info(f"本文フォルダが存在しませんが、作成可能: {honbun_folder}")
                return honbun_folder
            else:
                self.logger.error(f"親フォルダが存在しません: {ncode_folder}")
                return None
    
