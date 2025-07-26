"""Word譁・嶌蜃ｦ逅・Δ繧ｸ繝･繝ｼ繝ｫ"""
import zipfile
import tempfile
from pathlib import Path
from typing import List, Optional
import docx

from utils.logger import get_logger


class WordProcessor:
    """Word譁・嶌繧貞・逅・☆繧九け繝ｩ繧ｹ"""
    
    def __init__(self):
        """WordProcessor繧貞・譛溷喧"""
        self.logger = get_logger(__name__)
    
    def process_word_files(self, folder_path: Path) -> int:
        """
        繝輔か繝ｫ繝蜀・・縺吶∋縺ｦ縺ｮWord繝輔ぃ繧､繝ｫ繧貞・逅・        
        Args:
            folder_path: 蜃ｦ逅・☆繧九ヵ繧ｩ繝ｫ繝縺ｮ繝代せ
        
        Returns:
            蜃ｦ逅・＠縺溘ヵ繧｡繧､繝ｫ謨ｰ
        """
        self.logger.info(f"Word蜃ｦ逅・幕蟋・ {folder_path}")
        
        if not folder_path.exists():
            self.logger.error(f"繝輔か繝ｫ繝縺悟ｭ伜惠縺励∪縺帙ｓ: {folder_path}")
            return 0
        
        # .docx繝輔ぃ繧､繝ｫ繧呈､懃ｴ｢
        word_files = list(folder_path.glob("**/*.docx"))
        
        if not word_files:
            self.logger.warning("Word繝輔ぃ繧､繝ｫ縺瑚ｦ九▽縺九ｊ縺ｾ縺帙ｓ縺ｧ縺励◆")
            return 0
        
        processed_count = 0
        for word_file in word_files:
            try:
                if self.remove_first_line(word_file):
                    processed_count += 1
            except Exception as e:
                self.logger.error(f"繝輔ぃ繧､繝ｫ蜃ｦ逅・お繝ｩ繝ｼ {word_file}: {e}")
        
        self.logger.info(f"Word蜃ｦ逅・ｮ御ｺ・ {processed_count}/{len(word_files)} 繝輔ぃ繧､繝ｫ")
        return processed_count
    
    def remove_first_line(self, doc_path: Path) -> bool:
        """
        Word譁・嶌縺ｮ1陦檎岼繧貞炎髯､
        
        Args:
            doc_path: 蜃ｦ逅・☆繧妓ord譁・嶌縺ｮ繝代せ
        
        Returns:
            蜃ｦ逅・′謌仙粥縺励◆蝣ｴ蜷・rue
        """
        try:
            self.logger.info(f"1陦檎岼蜑企勁蜃ｦ逅・ {doc_path}")
            
            # 繝峨く繝･繝｡繝ｳ繝医ｒ髢九￥
            doc = docx.Document(doc_path)
            
            # 繝代Λ繧ｰ繝ｩ繝輔′蟄伜惠縺吶ｋ縺狗｢ｺ隱・            if not doc.paragraphs:
                self.logger.warning(f"遨ｺ縺ｮ繝峨く繝･繝｡繝ｳ繝・ {doc_path}")
                return True
            
            # 譛蛻昴・繝代Λ繧ｰ繝ｩ繝輔ｒ蜑企勁
            first_para = doc.paragraphs[0]
            p = first_para._element
            p.getparent().remove(p)
            
            # 菫晏ｭ・            doc.save(doc_path)
            
            self.logger.info(f"1陦檎岼繧貞炎髯､縺励∪縺励◆: {doc_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"1陦檎岼蜑企勁繧ｨ繝ｩ繝ｼ {doc_path}: {e}")
            return False
    
    def get_word_files(self, folder_path: Path) -> List[Path]:
        """
        繝輔か繝ｫ繝蜀・・Word繝輔ぃ繧､繝ｫ縺ｮ繝ｪ繧ｹ繝医ｒ蜿門ｾ・        
        Args:
            folder_path: 讀懃ｴ｢縺吶ｋ繝輔か繝ｫ繝縺ｮ繝代せ
        
        Returns:
            Word繝輔ぃ繧､繝ｫ縺ｮ繝代せ繝ｪ繧ｹ繝・        """
        if not folder_path.exists():
            return []
        
        return list(folder_path.glob("**/*.docx"))
    
    def process_zip_file(self, zip_path: Path, temp_dir: Optional[Path] = None) -> List[Path]:
        """
        ZIP繝輔ぃ繧､繝ｫ縺九ｉWord繝輔ぃ繧､繝ｫ繧呈歓蜃ｺ縺励・陦檎岼繧貞炎髯､
        
        Args:
            zip_path: 蜃ｦ逅・☆繧技IP繝輔ぃ繧､繝ｫ縺ｮ繝代せ
            temp_dir: 荳譎ゅョ繧｣繝ｬ繧ｯ繝医Μ・域欠螳壹＠縺ｪ縺・ｴ蜷医・閾ｪ蜍穂ｽ懈・・・        
        Returns:
            蜃ｦ逅・ｸ医∩Word繝輔ぃ繧､繝ｫ縺ｮ繝代せ繝ｪ繧ｹ繝・        """
        self.logger.info(f"ZIP蜃ｦ逅・幕蟋・ {zip_path}")
        
        if not zip_path.exists():
            self.logger.error(f"ZIP繝輔ぃ繧､繝ｫ縺悟ｭ伜惠縺励∪縺帙ｓ: {zip_path}")
            return []
        
        # 荳譎ゅョ繧｣繝ｬ繧ｯ繝医Μ縺ｮ貅門ｙ
        if temp_dir is None:
            temp_dir = Path(tempfile.mkdtemp())
        else:
            temp_dir.mkdir(parents=True, exist_ok=True)
        
        processed_files = []
        
        try:
            # ZIP繝輔ぃ繧､繝ｫ繧貞ｱ暮幕
            with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                zip_ref.extractall(temp_dir)
            
            self.logger.info(f"ZIP螻暮幕螳御ｺ・ {temp_dir}")
            
            # 螻暮幕縺輔ｌ縺欷ord繝輔ぃ繧､繝ｫ繧呈､懃ｴ｢
            word_files = self.get_word_files(temp_dir)
            
            if not word_files:
                self.logger.warning("Word繝輔ぃ繧､繝ｫ縺瑚ｦ九▽縺九ｊ縺ｾ縺帙ｓ縺ｧ縺励◆")
                return []
            
            self.logger.info(f"Word繝輔ぃ繧､繝ｫ {len(word_files)}蛟九ｒ逋ｺ隕・)
            
            # 蜷Цord繝輔ぃ繧､繝ｫ縺ｮ1陦檎岼繧貞炎髯､
            for word_file in word_files:
                try:
                    if self.remove_first_line(word_file):
                        processed_files.append(word_file)
                        self.logger.info(f"蜃ｦ逅・ｮ御ｺ・ {word_file.name}")
                    else:
                        self.logger.error(f"蜃ｦ逅・､ｱ謨・ {word_file.name}")
                except Exception as e:
                    self.logger.error(f"繝輔ぃ繧､繝ｫ蜃ｦ逅・お繝ｩ繝ｼ {word_file}: {e}")
            
            self.logger.info(f"ZIP蜃ｦ逅・ｮ御ｺ・ {len(processed_files)}/{len(word_files)} 繝輔ぃ繧､繝ｫ")
            return processed_files
            
        except zipfile.BadZipFile:
            self.logger.error(f"辟｡蜉ｹ縺ｪZIP繝輔ぃ繧､繝ｫ: {zip_path}")
            return []
        except Exception as e:
            self.logger.error(f"ZIP蜃ｦ逅・お繝ｩ繝ｼ: {e}")
            return []
    
    def find_ncode_folder(self, ncode: str) -> Optional[Path]:
        """
        謖・ｮ壹＆繧後◆N繧ｳ繝ｼ繝峨・繝輔か繝ｫ繝繧呈､懃ｴ｢
        
        Args:
            ncode: 讀懃ｴ｢縺吶ｋN繧ｳ繝ｼ繝・        
        Returns:
            隕九▽縺九▲縺櫻繝輔か繝ｫ繝縺ｮ繝代せ縲∬ｦ九▽縺九ｉ縺ｪ縺・ｴ蜷医・None
        """
        base_path = Path("G:/.shortcut-targets-by-id/0B6euJ_grVeOeMnJLU1IyUWgxeWM/NP-IRD")
        ncode_folder = base_path / ncode
        
        if ncode_folder.exists() and ncode_folder.is_dir():
            self.logger.info(f"N繝輔か繝ｫ繝逋ｺ隕・ {ncode_folder}")
            return ncode_folder
        else:
            self.logger.warning(f"N繝輔か繝ｫ繝縺瑚ｦ九▽縺九ｊ縺ｾ縺帙ｓ: {ncode_folder}")
            return None
    
    def find_honbun_folder(self, ncode_folder: Path) -> Optional[Path]:
        """
        N繝輔か繝ｫ繝驟堺ｸ九・縲梧悽譁・阪ヵ繧ｩ繝ｫ繝繧呈､懃ｴ｢
        
        Args:
            ncode_folder: N繝輔か繝ｫ繝縺ｮ繝代せ
        
        Returns:
            隕九▽縺九▲縺滓悽譁・ヵ繧ｩ繝ｫ繝縺ｮ繝代せ縲∬ｦ九▽縺九ｉ縺ｪ縺・ｴ蜷医・None
        """
        honbun_folder = ncode_folder / "譛ｬ譁・
        
        if honbun_folder.exists() and honbun_folder.is_dir():
            self.logger.info(f"譛ｬ譁・ヵ繧ｩ繝ｫ繝逋ｺ隕・ {honbun_folder}")
            return honbun_folder
        else:
            # 譛ｬ譁・ヵ繧ｩ繝ｫ繝縺悟ｭ伜惠縺励↑縺・ｴ蜷医・菴懈・蜿ｯ閭ｽ縺九メ繧ｧ繝・け
            if ncode_folder.exists():
                self.logger.info(f"譛ｬ譁・ヵ繧ｩ繝ｫ繝縺悟ｭ伜惠縺励∪縺帙ｓ縺後∽ｽ懈・蜿ｯ閭ｽ: {honbun_folder}")
                return honbun_folder
            else:
                self.logger.error(f"隕ｪ繝輔か繝ｫ繝縺悟ｭ伜惠縺励∪縺帙ｓ: {ncode_folder}")
                return None
    
